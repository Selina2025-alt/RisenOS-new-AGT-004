from __future__ import annotations

import hashlib
import json
import re
import shutil
import sys
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


WORKSPACE = Path(r"D:\company project\瑞森宣发智能体家族\RisenOS（新版）\AGT-RSN-004-Workspace")
OUTPUT_ROOT = WORKSPACE / "knowledge" / "sources" / "ingested"

SOURCES = [
    (
        "SRC-20260716-ALLHANDS",
        Path(r"D:\company project\0.BSC内容资产库\7月16日_老板全员大会演讲稿整理版_AItoB统一修订.docx"),
        "内部讲话与创始人叙事",
    ),
    (
        "SRC-20260718-COMPANY-PRODUCT",
        Path(r"D:\company project\0.BSC内容资产库\20260718公司简介与产品介绍（重点关注2条产品线，4大产品矩阵）.docx"),
        "企业介绍与产品矩阵",
    ),
    (
        "SRC-202607-WEBSITE",
        Path(r"D:\company project\0.BSC内容资产库\艾氪智能新官网设计 202607.docx"),
        "官网内容、信息架构与视觉需求",
    ),
    (
        "SRC-V8-AGENTIC-OS",
        Path(r"D:\company project\0.BSC内容资产库\通用资料库\01_艾氪智能产业级 Agentic OS 战略与壁垒手册V8.0.docx"),
        "产业级 Agentic OS 战略与壁垒",
    ),
    (
        "SRC-V7-GROUP-STRATEGY",
        Path(r"D:\company project\0.BSC内容资产库\通用资料库\01_艾氪智能集团战略手册 V7.0.docx"),
        "集团战略",
    ),
    (
        "SRC-20260724-LEADERSHIP-REPORT",
        Path(r"D:\company project\0.BSC内容资产库\通用资料库\01_《JovaAI 艾氪智能全球领先性研究报告：从企业智能体到产业级 Agentic OS 的代际跃迁》.docx"),
        "全球领先性研究与竞品参照",
    ),
]

TOPIC_KEYWORDS = {
    "企业战略": ["战略", "使命", "愿景", "集团战略", "转型", "商业模式", "产业AI", "第一性原理"],
    "品牌定位": ["品牌", "定位", "口号", "价值主张", "叙事", "关于我们", "企业简介"],
    "产品与架构": ["产品线", "产品矩阵", "JovaOS", "Agentic OS", "架构", "平台", "工作站", "Studio", "Wtree", "ICB"],
    "智能体能力": ["专业智能体", "智能体员工", "智能体团队", "多智能体", "协同", "上岗", "Agent Team"],
    "客户与场景": ["目标客户", "客户", "ICP", "为谁服务", "链主", "用户画像", "决策者", "供应链", "采购", "财务", "关务", "客服"],
    "案例与证据": ["案例", "客户案例", "数据", "证据", "验证", "领先", "全球首", "第一", "300", "3000"],
    "官网与视觉": ["官网", "页面", "首屏", "视觉", "颜色", "字体", "Logo", "图片", "导航", "交互"],
    "历史内容": ["演讲", "讲话", "文章", "内容", "金句", "故事", "发布会"],
    "竞品与壁垒": ["竞品", "竞争", "壁垒", "对比", "领先", "OpenAI", "Microsoft", "Salesforce", "SAP", "Oracle", "Dify"],
    "合规与保密": ["内部", "机密", "保密", "不得公开", "合规", "版权", "风险"],
}

HEADING_PATTERNS = [
    (1, re.compile(r"^第[一二三四五六七八九十百零〇0-9]+篇(?:\s|[:：]|$)")),
    (2, re.compile(r"^第[一二三四五六七八九十百零〇0-9]+章(?:\s|[:：]|$)")),
    (2, re.compile(r"^第\s*[0-9一二三四五六七八九十百零〇]+\s*[章部分](?:\s|[:：]|$)")),
    (2, re.compile(r"^[一二三四五六七八九十百]+[、．.]\s*")),
    (3, re.compile(r"^第[一二三四五六七八九十百零〇0-9]+节(?:\s|[:：]|$)")),
    (3, re.compile(r"^[（(][一二三四五六七八九十百0-9]+[）)]\s*")),
    (3, re.compile(r"^[0-9]{1,3}[、．.]\s*\S")),
    (2, re.compile(r"^(前言|序言|导言|摘要|执行摘要|结语|结论|附录)(?:\s|[:：]|$)")),
]


@dataclass
class Block:
    order: int
    kind: str
    locator: str
    text: str
    style: str = ""
    heading_level: int | None = None


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def safe_name(text: str, max_len: int = 70) -> str:
    text = re.sub(r"[\x00-\x1f<>:\"/\\|?*]", "_", text.strip())
    text = re.sub(r"\s+", "_", text)
    text = re.sub(r"_+", "_", text).strip("._")
    return (text or "未命名")[:max_len]


def paragraph_heading_level(paragraph: Paragraph) -> int | None:
    text = paragraph.text.strip()
    if not text:
        return None
    style = paragraph.style.name if paragraph.style else ""
    match = re.search(r"(?:Heading|标题)\s*([1-9])", style, re.IGNORECASE)
    if match:
        return int(match.group(1))
    if style.lower() == "title" or style == "标题":
        return 1
    if len(text) > 100 or text.endswith(("。", "；", ";")):
        return None
    for level, pattern in HEADING_PATTERNS:
        if pattern.search(text):
            return level
    runs = [run for run in paragraph.runs if run.text.strip()]
    if runs and len(text) <= 60:
        bold_ratio = sum(len(run.text) for run in runs if run.bold) / max(1, sum(len(run.text) for run in runs))
        sizes = [run.font.size.pt for run in runs if run.font.size]
        if bold_ratio >= 0.8 and sizes and max(sizes) >= 14:
            return 3
    return None


def iter_body_blocks(document: DocumentObject) -> Iterable[Paragraph | Table]:
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def table_to_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            value = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            value = re.sub(r"\s+", " ", value).replace("|", "\\|")
            cells.append(value)
        rows.append(cells)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
    return "\n".join(lines)


def extract_blocks(document: DocumentObject) -> list[Block]:
    blocks: list[Block] = []
    p_idx = 0
    t_idx = 0
    for order, item in enumerate(iter_body_blocks(document), start=1):
        if isinstance(item, Paragraph):
            p_idx += 1
            text = item.text.strip()
            if not text:
                continue
            blocks.append(
                Block(
                    order=order,
                    kind="paragraph",
                    locator=f"P{p_idx:05d}",
                    text=text,
                    style=item.style.name if item.style else "",
                    heading_level=paragraph_heading_level(item),
                )
            )
        else:
            t_idx += 1
            text = table_to_markdown(item)
            if text:
                blocks.append(
                    Block(
                        order=order,
                        kind="table",
                        locator=f"T{t_idx:04d}",
                        text=text,
                        style="Table",
                    )
                )
    return blocks


def score_topics(text: str) -> dict[str, int]:
    lowered = text.lower()
    scores = {}
    for topic, keywords in TOPIC_KEYWORDS.items():
        score = sum(lowered.count(word.lower()) for word in keywords)
        if score:
            scores[topic] = score
    return dict(sorted(scores.items(), key=lambda item: (-item[1], item[0])))


def split_sections(blocks: list[Block], max_chars: int = 18000) -> list[dict]:
    sections: list[dict] = []
    current: list[Block] = []
    current_title = "文档前置内容"
    current_level = 1
    current_chars = 0

    def flush() -> None:
        nonlocal current, current_chars
        if not current:
            return
        body = "\n\n".join(block.text for block in current)
        sections.append(
            {
                "title": current_title,
                "level": current_level,
                "start_locator": current[0].locator,
                "end_locator": current[-1].locator,
                "block_count": len(current),
                "char_count": len(body),
                "topics": score_topics(body),
                "blocks": current,
            }
        )
        current = []
        current_chars = 0

    for block in blocks:
        should_break = bool(block.heading_level and current and current_chars >= 800)
        if should_break:
            flush()
        if block.heading_level:
            current_title = block.text
            current_level = block.heading_level
        if current and current_chars + len(block.text) > max_chars:
            previous_title = current_title
            flush()
            current_title = f"{previous_title}（续）"
        current.append(block)
        current_chars += len(block.text)
    flush()
    return sections


def extract_package_media(source: Path, target: Path) -> list[dict]:
    target.mkdir(parents=True, exist_ok=True)
    manifest = []
    with zipfile.ZipFile(source) as archive:
        names = sorted(name for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/"))
        for index, name in enumerate(names, start=1):
            data = archive.read(name)
            original = Path(name).name
            output_name = f"{index:03d}_{safe_name(original)}"
            output = target / output_name
            output.write_bytes(data)
            manifest.append(
                {
                    "package_path": name,
                    "file": output_name,
                    "bytes": len(data),
                    "sha256": hashlib.sha256(data).hexdigest().upper(),
                }
            )
    return manifest


def header_footer_text(document: DocumentObject) -> dict:
    result = {"headers": [], "footers": []}
    for section_index, section in enumerate(document.sections, start=1):
        for kind, container in (("headers", section.header), ("footers", section.footer)):
            text = "\n".join(p.text.strip() for p in container.paragraphs if p.text.strip())
            if text:
                result[kind].append({"section": section_index, "text": text})
    return result


def write_source(source_id: str, source: Path, declared_use: str) -> dict:
    if not source.exists():
        raise FileNotFoundError(source)
    target = OUTPUT_ROOT / source_id
    if target.exists():
        shutil.rmtree(target)
    sections_dir = target / "sections"
    media_dir = target / "assets"
    sections_dir.mkdir(parents=True, exist_ok=True)

    document = Document(source)
    blocks = extract_blocks(document)
    sections = split_sections(blocks)
    media = extract_package_media(source, media_dir)
    headers_footers = header_footer_text(document)
    stat = source.stat()
    core = document.core_properties

    full_lines = [
        f"# {source.name}",
        "",
        f"- 来源ID：`{source_id}`",
        f"- 原始路径：`{source}`",
        f"- SHA-256：`{sha256_file(source)}`",
        f"- 声明用途：{declared_use}",
        "- 说明：以下内容按 Word 文档正文顺序完整提取；定位编号 P=正文段落，T=表格。",
        "",
    ]
    for block in blocks:
        prefix = "#" * min(6, block.heading_level) + " " if block.heading_level else ""
        full_lines.extend([f"<!-- {block.locator}; style={block.style} -->", prefix + block.text, ""])
    (target / "full_text.md").write_text("\n".join(full_lines), encoding="utf-8")

    section_index = []
    for index, section in enumerate(sections, start=1):
        filename = f"{index:03d}_{safe_name(section['title'])}.md"
        lines = [
            "---",
            f"source_id: {source_id}",
            f"source_file: {json.dumps(source.name, ensure_ascii=False)}",
            f"section_no: {index}",
            f"source_range: {section['start_locator']}-{section['end_locator']}",
            f"heading_level: {section['level']}",
            f"confidentiality: INTERNAL_SOURCE",
            "---",
            "",
            f"# {section['title']}",
            "",
            f"> 来源：{source.name}；定位：{section['start_locator']}–{section['end_locator']}。",
            "",
        ]
        for block in section["blocks"]:
            prefix = "#" * min(6, block.heading_level) + " " if block.heading_level else ""
            lines.extend([f"<!-- {block.locator}; style={block.style} -->", prefix + block.text, ""])
        (sections_dir / filename).write_text("\n".join(lines), encoding="utf-8")
        section_index.append(
            {
                key: value
                for key, value in section.items()
                if key != "blocks"
            }
            | {"file": f"sections/{filename}"}
        )

    manifest = {
        "source_id": source_id,
        "source_file": source.name,
        "original_path": str(source),
        "declared_use": declared_use,
        "sha256": sha256_file(source),
        "size_bytes": stat.st_size,
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        "ingested_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "core_properties": {
            "title": core.title,
            "subject": core.subject,
            "author": core.author,
            "keywords": core.keywords,
            "comments": core.comments,
            "created": core.created.isoformat() if core.created else None,
            "modified": core.modified.isoformat() if core.modified else None,
            "last_modified_by": core.last_modified_by,
            "revision": core.revision,
        },
        "counts": {
            "document_paragraphs": len(document.paragraphs),
            "nonempty_document_paragraphs": sum(bool(p.text.strip()) for p in document.paragraphs),
            "tables": len(document.tables),
            "table_rows": sum(len(table.rows) for table in document.tables),
            "ordered_nonempty_blocks": len(blocks),
            "ordered_paragraph_blocks": sum(block.kind == "paragraph" for block in blocks),
            "ordered_table_blocks": sum(block.kind == "table" for block in blocks),
            "sections": len(sections),
            "media_assets": len(media),
            "extracted_characters": sum(len(block.text) for block in blocks),
        },
        "headers_footers": headers_footers,
        "media": media,
        "section_index": section_index,
    }
    (target / "source_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return manifest


def main() -> int:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    manifests = []
    for source_id, source, declared_use in SOURCES:
        print(f"Ingesting {source_id}: {source.name}", flush=True)
        manifests.append(write_source(source_id, source, declared_use))
    batch = {
        "batch_id": "BATCH-20260730-BSC-6DOCX",
        "ingested_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "source_count": len(manifests),
        "total_blocks": sum(item["counts"]["ordered_nonempty_blocks"] for item in manifests),
        "total_characters": sum(item["counts"]["extracted_characters"] for item in manifests),
        "total_sections": sum(item["counts"]["sections"] for item in manifests),
        "total_media_assets": sum(item["counts"]["media_assets"] for item in manifests),
        "sources": [
            {
                "source_id": item["source_id"],
                "source_file": item["source_file"],
                "sha256": item["sha256"],
                "counts": item["counts"],
            }
            for item in manifests
        ],
    }
    (OUTPUT_ROOT / "batch_manifest.json").write_text(
        json.dumps(batch, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(batch, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
