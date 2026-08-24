from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = "17365D"
LIGHT_BLUE = "DCE6F1"
GRAY = "666666"


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9)


def clean_inline(text: str) -> str:
    cleaned = (
        text.replace("`", "")
        .replace("**", "")
        .replace("<br>", " ")
        .replace("<br/>", " ")
    )
    return re.sub(r"\*([^*]+)\*", r"\1", cleaned)


def configure_document(
    document: Document,
    header_text: str = "艾氪智能｜AGT-RSN-004 内容合规知识库",
) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.35

    for name, size, color in (
        ("Title", 22, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, "1F4E79"),
        ("Heading 4", 10, "365F91"),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = header_text
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, 9)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = clean_inline(row[column_index] if column_index < len(row) else "")
            if row_index == 0:
                shade_cell(cell, BLUE)
            elif row_index % 2 == 0:
                shade_cell(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    set_run_font(run, 8.5, row_index == 0)
                    if row_index == 0:
                        run.font.color.rgb = RGBColor(255, 255, 255)
    document.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        index += 1
    return rows, index


def convert(
    source: Path,
    target: Path,
    header_text: str = "艾氪智能｜AGT-RSN-004 内容合规知识库",
) -> None:
    document = Document()
    configure_document(document, header_text=header_text)
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    title_seen = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1 and not title_seen:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                set_run_font(run, 22, True)
                title_seen = True
            else:
                document.add_heading(text, level=level)
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            paragraph.paragraph_format.right_indent = Cm(0.5)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(clean_inline(stripped.lstrip("> ")))
            set_run_font(run, 9.5)
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(GRAY)
            index += 1
            continue

        checkbox = re.match(r"^-\s+\[([ xX])\]\s+(.+)$", stripped)
        if checkbox:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(
                f"{'☑' if checkbox.group(1).lower() == 'x' else '☐'} "
                f"{clean_inline(checkbox.group(2))}"
            )
            set_run_font(run, 10)
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(clean_inline(bullet.group(1)))
            set_run_font(run, 10)
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            run = paragraph.add_run(clean_inline(numbered.group(1)))
            set_run_font(run, 10)
            index += 1
            continue

        paragraph = document.add_paragraph()
        run = paragraph.add_run(clean_inline(stripped))
        set_run_font(run, 10.5)
        index += 1

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


if __name__ == "__main__":
    if len(sys.argv) not in {3, 4}:
        raise SystemExit("usage: markdown_to_docx.py SOURCE.md TARGET.docx [HEADER_TEXT]")
    convert(
        Path(sys.argv[1]),
        Path(sys.argv[2]),
        header_text=sys.argv[3] if len(sys.argv) == 4 else "艾氪智能｜AGT-RSN-004 内容合规知识库",
    )
